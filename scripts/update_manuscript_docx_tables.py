"""
Update compact-object data tables in the manuscript Word doc to SI units
and pipeline-computed values (WD eta=0.007, SMBH eta=0.057, NS/BH tabulated ERD).
"""

from __future__ import annotations

import sys
from pathlib import Path

import pandas as pd
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from physics_engine import select_plot_compact_results  # noqa: E402

DEFAULT_INPUT = Path(
    r"C:\Users\LuchK\Downloads\Vidal Kadian 2026 - Power Density as a Universal Biosignature.docx"
)
DEFAULT_OUTPUT = DEFAULT_INPUT.with_name(
    DEFAULT_INPUT.stem + " - SI tables updated.docx"
)

HEADERS = (
    "Name",
    "Mass (kg)",
    "Accretion rate (kg s-1)",
    "Power density Phi_m (W kg-1)",
)

TABLE_SPECS: tuple[tuple[int, str], ...] = (
    (0, "cataclysmic_variables"),
    (1, "neutron_stars"),
    (2, "transient_black_holes"),
    (3, "supermassive_black_holes"),
)


def format_sci(value: float) -> str:
    """European-style scientific notation (comma decimal separator)."""
    text = f"{float(value):.4E}"
    return text.replace(".", ",")


def format_phi(value: float) -> str:
    value = float(value)
    if 1.0e-4 <= abs(value) < 1.0e4:
        text = f"{value:.4g}"
    else:
        text = f"{value:.4E}"
    return text.replace(".", ",")


def set_row(cells, name: str, mass_kg: float, mdot_kg_s: float, phi_w_kg: float) -> None:
    cells[0].text = name
    cells[1].text = format_sci(mass_kg)
    cells[2].text = format_sci(mdot_kg_s)
    cells[3].text = format_phi(phi_w_kg)


def is_summary_row(name: str) -> bool:
    return name.strip().upper() in {"AVERAGE", "MEAN"}


def update_table(
    table,
    frame: pd.DataFrame,
) -> tuple[int, int]:
    """Update header + data rows; recompute trailing AVERAGE row."""
    for col, header in enumerate(HEADERS):
        table.rows[0].cells[col].text = header

    by_name = frame.set_index("name")
    updated = 0
    masses: list[float] = []
    mdots: list[float] = []
    phis: list[float] = []

    for row_idx in range(1, len(table.rows)):
        name = table.rows[row_idx].cells[0].text.strip()
        if not name or is_summary_row(name):
            continue
        if name not in by_name.index:
            raise KeyError(f"Object '{name}' missing from processed CSV.")
        record = by_name.loc[name]
        if isinstance(record, pd.DataFrame):
            record = record.iloc[0]
        mass_kg = float(record["mass_kg"])
        mdot_kg_s = float(record["mdot_kg_s"])
        phi_w_kg = float(record["power_density_w_per_kg"])
        set_row(table.rows[row_idx].cells, name, mass_kg, mdot_kg_s, phi_w_kg)
        masses.append(mass_kg)
        mdots.append(mdot_kg_s)
        phis.append(phi_w_kg)
        updated += 1

    if table.rows:
        last_name = table.rows[-1].cells[0].text.strip()
        if is_summary_row(last_name) and masses:
            set_row(
                table.rows[-1].cells,
                last_name,
                sum(masses) / len(masses),
                sum(mdots) / len(mdots),
                sum(phis) / len(phis),
            )

    return updated, len(masses)


def main(input_path: Path = DEFAULT_INPUT, output_path: Path = DEFAULT_OUTPUT) -> None:
    if not input_path.exists():
        raise FileNotFoundError(input_path)

    compact = pd.read_csv(ROOT / "processed" / "processed_compact_results.csv")
    smbh = pd.read_csv(ROOT / "processed" / "processed_smbh_results.csv")
    plot_compact = select_plot_compact_results(compact)

    doc = Document(str(input_path))
    if len(doc.tables) < 4:
        raise ValueError(f"Expected at least 4 tables, found {len(doc.tables)}")

    for table_idx, category in TABLE_SPECS:
        if category == "supermassive_black_holes":
            frame = smbh[smbh["track"] == "gravitational"].copy()
        else:
            frame = plot_compact[plot_compact["category"] == category].copy()
        n_updated, n_data = update_table(doc.tables[table_idx], frame)
        print(f"Table {table_idx + 1} ({category}): updated {n_updated}/{n_data} rows")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    inp = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_INPUT
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else DEFAULT_OUTPUT
    main(inp, out)
