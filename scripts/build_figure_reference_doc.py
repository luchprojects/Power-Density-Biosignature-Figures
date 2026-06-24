"""
Generate a concise PI-facing figure reference (Word).

One section per PDF: purpose, data source, equations, and brief notes only.
No color tables, code paths, or audit inventories — those live on the figures / in docs/FIGURES.md.
"""

from __future__ import annotations

import sys
from datetime import datetime, timezone
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

import config

DOC_PATH = config.FIGURE_REFERENCE_DOC

# Shared reference-overlay blurb (biology + unified figures only).
CHAISSON_OVERLAY_NOTE = (
    "Literature overlays (not data): Chaisson (2001) modern society benchmark "
    f"(~{config.CHAISSON_2001_SOCIETY_W_PER_KG:g} W·kg⁻¹); "
    f"van Duin (2024) stability limit at 10⁵ W·kg⁻¹."
)


def _add_field(doc, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.add_run(f"{label}. ").bold = True
    p.add_run(text)


def _figure_entries(
    *,
    yso_n: int,
    cv_n: int,
    ns_n: int,
    bh_n: int,
    bio_n: int,
    smbH_n: int,
) -> list[tuple[str, dict[str, str]]]:
    return [
        (
            "figure_unified_master.pdf",
            {
                "Purpose": (
                    "Overview of the full power-density continuum — biology, young stellar objects "
                    "(YSOs), and accreting compact objects on one log–log plot."
                ),
                "Data source": (
                    f"Biology — van Duin et al. (2024) MOESM1 Section I "
                    f"(data/biology/von_duin_2024_erd_moesm1.csv; {bio_n} measurements). "
                    f"YSOs — Manara et al. (2022) PPVII compilation "
                    f"(data/yso/mdots_forclement.dat; {yso_n} systems; Alcalá+2017 & Manara+2017 filter; "
                    "Somers 2020 SPOTS masses). "
                    f"Compact objects — Vidal (2020) tables "
                    f"(data/compact/Power density data.csv; {cv_n} white dwarfs [Dubus et al. 2018], "
                    f"{ns_n} neutron stars [Galloway et al. 2008], {bh_n} black holes [Coriat et al. 2012])."
                ),
                "Equations": (
                    "All panels use Φ_m = L / M (W·kg⁻¹, mass in kg). "
                    "Biology: Φ_m from tabulated empirical ERD. "
                    "YSOs: Φ_m = L_acc / M★. "
                    "Compact objects: η = GM/(Rc²), L = η Ṁ c², Φ_m = L/M; tabulated ERD used when available."
                ),
                "Notes": (
                    f"{CHAISSON_OVERLAY_NOTE} WDs: green = Dubus Table A.2 (general CVs), "
                    "fuchsia = Table A.3 (nova-like); YSO: open yellow rings on top. "
                    "No error bars on this panel — see figure_wd_dubus_uncertainties.pdf."
                ),
            },
        ),
        (
            "figure_biology.pdf",
            {
                "Purpose": (
                    "Zoom on biological systems — empirical energy rate density (ERD) measurements "
                    "grouped by prokaryotes, unicellular eukaryotes, and multicellular organisms."
                ),
                "Data source": (
                    f"van Duin et al. (2024) MOESM1 Section I — same table as unified "
                    f"({bio_n} measurements; processed copy in processed/processed_von_duin_biology.csv)."
                ),
                "Equations": (
                    "Φ_m taken directly from the published ERD column (already in W·kg⁻¹)."
                ),
                "Notes": CHAISSON_OVERLAY_NOTE,
            },
        ),
        (
            "figure_yso.pdf",
            {
                "Purpose": (
                    "Abiotic control sample — accreting young stars only, without biology or compact objects."
                ),
                "Data source": (
                    f"Manara et al. (2022) PPVII → mdots_forclement.dat ({yso_n} systems). "
                    "Stellar masses from Baraffe+2015 with Somers (2020) SPOTS correction (17% spot coverage)."
                ),
                "Equations": (
                    "Φ_m = L_acc / M★, with L_acc = 10^(log L_acc) × L☉."
                ),
                "Notes": "Empirical scatter only; no literature reference overlays.",
            },
        ),
        (
            "figure_compact_objects.pdf",
            {
                "Purpose": (
                    "Accreting compact objects only — cataclysmic variables (accreting white dwarfs), "
                    "neutron stars, and transient black holes."
                ),
                "Data source": (
                    f"data/compact/Power density data.csv — {cv_n} WDs (Dubus et al. 2018), "
                    f"{ns_n} NSs (Galloway et al. 2008), {bh_n} BHs (Coriat et al. 2012)."
                ),
                "Equations": (
                    "η_grav = GM/(Rc²); L = η Ṁ c²; Φ_m = L/M. Tabulated ERD overrides the computed value when both are present."
                ),
                "Notes": (
                    "WD colors: green = Dubus Table A.2 (general CVs), fuchsia = Table A.3 (nova-like). "
                    "No YSO overlay. van Duin stability line shown; no Chaisson overlays."
                ),
            },
        ),
        (
            "figure_wd_dubus_uncertainties.pdf",
            {
                "Purpose": (
                    "Supplementary panel — white-dwarf systems with published uncertainty bars from "
                    "Dubus et al. (2018), kept separate from the main continuum figures."
                ),
                "Data source": (
                    f"WD positions from the compact-object table; uncertainties from "
                    f"Dubus et al. (2018) Tables A.2–A.3 "
                    f"(data/compact/dubus_2018_wd_uncertainties.csv; {cv_n} systems)."
                ),
                "Equations": (
                    "Horizontal errors: primary-star mass M₁ ± σ (Dubus). "
                    "Vertical errors: asymmetric 68% Monte Carlo interval on Ṁ, propagated to Φ_m "
                    "at the same accretion efficiency as each plotted point."
                ),
                "Notes": (
                    "Error bars only on this figure. Green = Dubus Table A.2; fuchsia = Table A.3. "
                    "No literature reference overlays."
                ),
            },
        ),
        (
            "figure_smbh_seyfert1.pdf",
            {
                "Purpose": (
                    "Test panel for Seyfert 1 supermassive black holes — separate from the unified "
                    "master continuum."
                ),
                "Data source": (
                    f"Vidal (2020) Table 5 — {smbH_n} Seyfert 1 SMBHs from Meyer-Hofmeister & Meyer (2011), "
                    "parsed from references/Vidal-2020 PDF into "
                    "data/compact/vidal_2020_table5_smbh_seyfert1.csv."
                ),
                "Equations": (
                    "Tabulated ERD (erg·s⁻¹·g⁻¹) at η = 0.1 per Vidal; Φ_m = L/M in W·kg⁻¹. "
                    "Schwarzschild radius used for the gravitational-efficiency track when computing "
                    "from Ṁ; tabulated ERD overrides when both are present."
                ),
                "Notes": (
                    "Not plotted on figure_unified_master.pdf. Chaisson modern society benchmark shown; "
                    "no error bars (none in source table)."
                ),
            },
        ),
    ]


def build_document(
    *,
    yso_n: int = 0,
    cv_n: int = 0,
    ns_n: int = 0,
    bh_n: int = 0,
    bio_n: int = 0,
    smbH_n: int = 0,
) -> Path:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    doc.add_heading("Figure Reference — Power Density as a Biosignature", level=0)
    doc.add_paragraph(
        f"Updated {datetime.now(timezone.utc).strftime('%d %B %Y')}. "
        "Brief guide to each PDF in figures/: what it shows, where the data come from, "
        "and the equations used. Colors and markers are defined in the figure legends."
    )

    doc.add_heading("Regenerating the figures", level=1)
    doc.add_paragraph(
        "Open a terminal in the project folder and run:  python main.py\n"
        "Requires: pip install -r requirements.txt (once). "
        "Outputs: six PDFs in figures/ and processed tables in processed/."
    )

    doc.add_heading("Quantity plotted", level=1)
    doc.add_paragraph(
        "Every figure is a log–log plot of mass (kg) versus mass-specific power density "
        "Φ_m = L / M in W·kg⁻¹. Each point is computed from a published table — nothing is placed by hand."
    )

    doc.add_heading("Figures", level=1)
    for filename, fields in _figure_entries(
        yso_n=yso_n,
        cv_n=cv_n,
        ns_n=ns_n,
        bh_n=bh_n,
        bio_n=bio_n,
        smbH_n=smbH_n,
    ):
        doc.add_heading(filename, level=2)
        for label, text in fields.items():
            _add_field(doc, label, text)
        doc.add_paragraph()

    doc.add_paragraph(
        "For file paths, pipeline details, and code entry points, see README.md and docs/FIGURES.md "
        "in the project folder."
    )

    DOC_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOC_PATH)
    return DOC_PATH


def build_document_from_pipeline(
    compact_results,
    yso_results,
    biology_results=None,
    smbH_results=None,
) -> Path:
    """Build the Word document using live pipeline DataFrame counts."""
    grav = compact_results[compact_results["track"] == "gravitational"]
    bio_n = 0
    if biology_results is not None and not biology_results.empty:
        bio_n = int(len(biology_results))
    smbH_n = 0
    if smbH_results is not None and not smbH_results.empty:
        smbH_n = int(len(smbH_results[smbH_results["track"] == "gravitational"]))

    return build_document(
        yso_n=len(yso_results),
        cv_n=int((grav["category"] == config.CATEGORY_CATACLYSMIC_VARIABLES).sum()),
        ns_n=int((grav["category"] == config.CATEGORY_NEUTRON_STARS).sum()),
        bh_n=int((grav["category"] == config.CATEGORY_TRANSIENT_BLACK_HOLES).sum()),
        bio_n=bio_n,
        smbH_n=smbH_n,
    )


if __name__ == "__main__":
    path = build_document()
    print(f"Wrote {path}")
